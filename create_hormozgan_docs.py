#!/usr/bin/env python3
# create_hormozgan_docs.py
# ساخته شده برای Termux: تولید Hormozgan_Driver_Pro_Final.docx و Hormozgan_Driver_Pro_Final.pdf

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
import os
import datetime

# ----------  متن مستند (خلاصه و منظم) ----------
# این متن فارسی‌محور است و ترجمه‌ی هر عنوان به انگلیسی بلافاصله زیرش می‌آید.
# (می‌توانی متن را در همین متغیرها ویرایش کنی قبل از اجرا)

title_fa = "سامانه هوشمند مدیریت سفر و گردشگران استان هرمزگان\nHormozgan Driver Pro"
author_line = "تهیه و توسعه توسط: مهدی حیدری پوری\nPrepared & Developed by: Mahdi Heidari Pouri"
generated_time = "تاریخ تولید / Generated: " + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

sections = [
    {
        "title_fa": "معرفی پروژه",
        "title_en": "Project Overview",
        "body_fa": (
            "Hormozgan Driver Pro یک سامانهٔ بومی هوشمند برای مدیریت سفرها، رانندگان و گردشگری استان هرمزگان است. "
            "هدف: افزایش امنیت، بهبود مدیریت ترافیک، و حمایت از گردشگری محلی."
        ),
        "body_en": (
            "Hormozgan Driver Pro is a localized smart platform for trip and driver management and tourism in Hormozgan Province. "
            "Goal: improve safety, traffic management, and support local tourism."
        )
    },
    {
        "title_fa": "معماری فنی",
        "title_en": "Technical Architecture",
        "body_fa": (
            "معماری: Monorepo با Node.js/Express، پایگاه‌داده Postgres + PostGIS، نقشه با Leaflet.js، و اپ موبایل React Native. "
            "قابلیت اجرا در محیط Termux و میزبانی در Render یا سرور داخلی."
        ),
        "body_en": (
            "Architecture: Monorepo using Node.js/Express, Postgres + PostGIS, Leaflet.js for mapping, and React Native mobile app. "
            "Can run in Termux and be hosted on Render or local servers."
        )
    },
    {
        "title_fa": "مأموریت گردشگری و توریسم",
        "title_en": "Tourism & Hospitality Mission",
        "body_fa": (
            "پشتیبانی از تور لیدرها، معرفی جاذبه‌ها، نمایش رویدادها و جشن‌های محلی، ارائه تخفیف‌های همکاری با هتل‌ها و رستوران‌ها، "
            "و ایجاد تجربهٔ میزبانی برای گردشگران."
        ),
        "body_en": (
            "Support tour guides, highlight attractions, show local events and festivals, offer collaborative discounts with hotels and restaurants, "
            "and create hospitality experiences for tourists."
        )
    },
    {
        "title_fa": "ماژول هوش مصنوعی",
        "title_en": "AI Smart Assistant Module",
        "body_fa": (
            "تشخیص ساعات اوج ترافیک، تحلیل داده‌های موقعیتی، پیش‌بینی بار ترافیکی، توصیه مسیر بهینه، "
            "و اطلاع‌رسانی لحظه‌ای دربارهٔ آب‌وهوا و هشدار سرعت."
        ),
        "body_en": (
            "Detect peak traffic hours, analyze location data, predict traffic loads, recommend optimized routes, "
            "and provide real-time weather and speed warnings."
        )
    },
    {
        "title_fa": "سامانه تماس اضطراری (SOS)",
        "title_en": "Emergency & SOS System",
        "body_fa": (
            "دکمه SOS با ارسال موقعیت GPS لحظه‌ای به مرکز امداد، تماس خودکار با اورژانس/پلیس/آتش‌نشانی، "
            "و امکان ارسال پیامک در صورت قطع اینترنت."
        ),
        "body_en": (
            "SOS button sends real-time GPS location to emergency center, auto-calls ambulance/police/fire, "
            "and supports SMS fallback if internet is offline."
        )
    },
    {
        "title_fa": "امنیت و احراز هویت بیومتریک",
        "title_en": "Security & Biometric Authentication",
        "body_fa": (
            "احراز هویت چندمرحله‌ای شامل OTP، تشخیص چهره و اثر انگشت، تطبیق با ثبت‌احوال و استعلام از پلیس راهور (پلیس‌من). "
            "داده‌ها رمزنگاری‌شده (TLS و AES-256) و دسترسی‌ها سطح‌بندی شده است."
        ),
        "body_en": (
            "Multi-factor authentication (OTP, face, fingerprint), verification with National Registry and Rahvar (police). "
            "Data encrypted via TLS and AES-256, with role-based access control."
        )
    },
    {
        "title_fa": "بخش رقابتی، جذب و تبلیغات محلی",
        "title_en": "Competitive Strategy, Acquisition & Local Ads",
        "body_fa": (
            "سیستم امتیازدهی رانندگان، تخفیف‌های هوشمند برای مسافران، تبلیغات جغرافیایی محلی (Geo Ads)، و مشارکت با کسب‌وکارهای محلی. "
            "هدف: ماندگاری کاربران و مزیت رقابتی نسبت به سرویس‌های غیربومی."
        ),
        "body_en": (
            "Driver ranking, smart promos for passengers, geo-targeted local ads, and partnerships with local businesses. "
            "Aim: retain users and compete against non-local platforms."
        )
    },
    {
        "title_fa": "تحلیل بازار و اثر اجتماعی",
        "title_en": "Market Analysis & Social Impact",
        "body_fa": (
            "افزایش اشتغال محلی، جذب گردشگر، ارتقای ایمنی و افزایش درآمد رانندگان محلی؛ "
            "پتانسیل توسعه به عنوان پلتفرم پایلوت شهر هوشمند."
        ),
        "body_en": (
            "Increase local employment, attract tourists, improve safety and drivers' income; potential to scale as a smart city pilot."
        )
    },
    {
        "title_fa": "نقشه راه توسعه",
        "title_en": "Roadmap",
        "body_fa": (
            "فاز 1: پیاده‌سازی هسته و دمو؛ فاز 2: اتصال به نهادهای رسمی و امنیت بیومتریک؛ "
            "فاز 3: توسعه اپ موبایل و ماژول توریسم؛ فاز 4: هوش پیش‌بینی و توسعه ملی."
        ),
        "body_en": (
            "Phase 1: core & demo; Phase 2: integration with official agencies & biometrics; "
            "Phase 3: mobile app & tourism module; Phase 4: predictive AI & national scaling."
        )
    },
    {
        "title_fa": "پیوست‌ها",
        "title_en": "Attachments",
        "body_fa": (
            "فایل‌های پیوست: سورس دمو، ویدیوی معرفی، مستندات ماژول موزیک، ماژول روزهای جهانی و فایل ساختار پروژه."
        ),
        "body_en": (
            "Attachments: demo source, intro video, music module doc, world days feature, project structure file."
        )
    },
    {
        "title_fa": "تهیه‌کننده",
        "title_en": "Prepared By",
        "body_fa": "مهدی حیدری پوری",
        "body_en": "Mahdi Heidari Pouri"
    }
]

# ----------  توابع برای ساخت DOCX و PDF ----------
def make_docx(filename="Hormozgan_Driver_Pro_Final.docx"):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Tahoma'
    style.font.size = Pt(12)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title_fa + "\n")
    run.bold = True
    run.font.size = Pt(16)
    p.add_run(author_line + "\n" + generated_time)

    doc.add_paragraph("")  # spacer

    # sections
    for s in sections:
        h = doc.add_paragraph()
        h.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = h.add_run(s["title_fa"] + "\n")
        run.bold = True
        run.font.size = Pt(14)

        # English subtitle (left-aligned)
        he = doc.add_paragraph()
        he.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run_e = he.add_run(s["title_en"] + "\n")
        run_e.italic = True
        run_e.font.size = Pt(12)

        # Persian body (right aligned)
        bf = doc.add_paragraph()
        bf.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        bf.add_run(s["body_fa"] + "\n")

        # English body (left)
        be = doc.add_paragraph()
        be.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        be.add_run(s["body_en"] + "\n")

        doc.add_paragraph("")

    # footer / prepared by
    doc.add_paragraph("")
    f = doc.add_paragraph()
    f.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    f.add_run("Prepared & Developed by / تهیه و توسعه توسط:\n" + "مهدی حیدری پوری").bold = True

    doc.save(filename)
    return filename

def make_pdf(filename="Hormozgan_Driver_Pro_Final.pdf"):
    doc = SimpleDocTemplate(filename, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    # define a style supporting UTF-8
    styles.add(ParagraphStyle(name='FA', fontName='Helvetica', fontSize=12, leading=14))
    story = []

    story.append(Paragraph(title_fa.replace("\n", "<br/>"), styles['Title']))
    story.append(Paragraph(author_line + "<br/>" + generated_time, styles['Normal']))
    story.append(Spacer(1, 6))

    for s in sections:
        story.append(Paragraph("<b>" + s["title_fa"] + "</b>", styles['Heading2']))
        story.append(Paragraph("<i>" + s["title_en"] + "</i>", styles['Normal']))
        story.append(Paragraph(s["body_fa"], styles['Normal']))
        story.append(Paragraph(s["body_en"], styles['Normal']))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Prepared & Developed by: مهدی حیدری پوری", styles['Normal']))

    doc.build(story)
    return filename

def main():
    out_docx = make_docx()
    out_pdf = make_pdf()
    # move to shared Download so user can easily access
    try:
        download_dir = "/sdcard/Download"
        if os.path.exists(download_dir):
            os.replace(out_docx, os.path.join(download_dir, out_docx))
            os.replace(out_pdf, os.path.join(download_dir, out_pdf))
            print("Files created and moved to /sdcard/Download:")
            print(os.path.join(download_dir, out_docx))
            print(os.path.join(download_dir, out_pdf))
        else:
            print("Files created in current directory:")
            print(os.path.abspath(out_docx))
            print(os.path.abspath(out_pdf))
    except Exception as e:
        print("Created files but could not move to /sdcard/Download. See current directory.")
        print("Error:", e)

if __name__ == "__main__":
    main()
